"""
Report export utilities.

Converts a persisted report (synthesis, comparison, or evaluation) into
a downloadable Markdown, PDF, or DOCX file - the format an analyst would
actually want to forward, not just a JSON blob.
"""

from __future__ import annotations

import io


def report_to_markdown(report_type: str, company: str, fiscal_year: str | None, payload: dict) -> str:
    lines = [f"# {report_type.title()} Report — {company}"]
    if fiscal_year:
        lines.append(f"**Fiscal Year:** {fiscal_year}\n")

    if report_type == "synthesis":
        lines.append(f"## Executive Summary\n\n{payload.get('executive_summary', '')}\n")
        lines.append("## Top Risks\n")
        for risk in payload.get("top_risks", []):
            lines.append(f"- **[{risk['severity'].upper()}]** ({risk['category']}) {risk['headline']}")
            lines.append(f"  {risk['detail']}")
        lines.append(f"\n## Recommendation\n\n{payload.get('recommendation', '')}\n")

    elif report_type == "comparison":
        lines.append(
            f"**Comparing:** {payload.get('prior_fiscal_year')} vs {payload.get('current_fiscal_year')}\n"
        )
        lines.append(f"## Overall Trajectory: {payload.get('overall_trajectory', '').upper()}\n")
        lines.append("## Material Trend Shifts\n")
        for shift in payload.get("trend_shifts", []):
            lines.append(
                f"- **[{shift['materiality'].upper()}]** {shift['metric']}: "
                f"{shift['prior_value']} -> {shift['current_value']} ({shift['direction']})"
            )
            lines.append(f"  {shift['note']}")
        lines.append(f"\n## Narrative\n\n{payload.get('narrative', '')}\n")

    elif report_type == "evaluation":
        lines.append(f"**Overall Accuracy:** {payload.get('overall_accuracy', 0) * 100:.1f}%\n")
        lines.append("## Results by Filing\n")
        for result in payload.get("results", []):
            lines.append(
                f"- {result.get('filing', 'unknown')}: "
                f"{result.get('passed_count')}/{result.get('total_count')} checks passed"
            )

    return "\n".join(lines)


def report_to_pdf_bytes(report_type: str, company: str, fiscal_year: str | None, payload: dict) -> bytes:
    from fpdf import FPDF

    markdown = report_to_markdown(report_type, company, fiscal_year, payload)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for line in markdown.split("\n"):
        clean_line = line.replace("**", "").replace("##", "").replace("#", "").strip()
        if not clean_line:
            pdf.ln(4)
            continue
        # Encode to latin-1 with replacement since core FPDF fonts don't support
        # full Unicode - acceptable for this report's ASCII-heavy financial text.
        safe_line = clean_line.encode("latin-1", errors="replace").decode("latin-1")
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, safe_line)

    output = pdf.output()
    return bytes(output)


def report_to_docx_bytes(report_type: str, company: str, fiscal_year: str | None, payload: dict) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(f"{report_type.title()} Report — {company}", level=1)
    if fiscal_year:
        doc.add_paragraph(f"Fiscal Year: {fiscal_year}")

    if report_type == "synthesis":
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(payload.get("executive_summary", ""))
        doc.add_heading("Top Risks", level=2)
        for risk in payload.get("top_risks", []):
            doc.add_paragraph(
                f"[{risk['severity'].upper()}] ({risk['category']}) {risk['headline']}: {risk['detail']}",
                style="List Bullet",
            )
        doc.add_heading("Recommendation", level=2)
        doc.add_paragraph(payload.get("recommendation", ""))

    elif report_type == "comparison":
        doc.add_paragraph(
            f"Comparing {payload.get('prior_fiscal_year')} vs {payload.get('current_fiscal_year')}"
        )
        doc.add_heading(f"Overall Trajectory: {payload.get('overall_trajectory', '').upper()}", level=2)
        doc.add_heading("Material Trend Shifts", level=2)
        for shift in payload.get("trend_shifts", []):
            doc.add_paragraph(
                f"[{shift['materiality'].upper()}] {shift['metric']}: "
                f"{shift['prior_value']} -> {shift['current_value']} "
                f"({shift['direction']}) - {shift['note']}",
                style="List Bullet",
            )
        doc.add_heading("Narrative", level=2)
        doc.add_paragraph(payload.get("narrative", ""))

    elif report_type == "evaluation":
        doc.add_paragraph(f"Overall Accuracy: {payload.get('overall_accuracy', 0) * 100:.1f}%")
        doc.add_heading("Results by Filing", level=2)
        for result in payload.get("results", []):
            doc.add_paragraph(
                f"{result.get('filing', 'unknown')}: "
                f"{result.get('passed_count')}/{result.get('total_count')} checks passed",
                style="List Bullet",
            )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
